import io
import os
import builtins
import streamlit as st
from docx import Document
from orchestrator import Orchestrator
from agents.blog_post_agent import BlogPostAgent
from agents.transcript_agent import TranscriptAgent
from database import init_db, get_video_by_id, save_video, update_blog_post

# ── Database init (safe via CREATE TABLE IF NOT EXISTS) ───────────────────────
try:
    init_db()
except Exception as e:
    st.warning(f"DB init failed: {e}")


def blog_post_to_docx(markdown_text: str) -> bytes:
    """Convert a markdown blog post string into a .docx file in memory."""
    try:
        doc = Document()
        for line in markdown_text.splitlines():
            stripped = line.strip()
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(stripped)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        print(f"[blog_post_to_docx] Failed to generate Word document: {e}")
        return b""

st.set_page_config(
    page_title="YouTube to Blog Post",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded",
)

with st.sidebar:
    st.title("📝 YT Blog Agent")
    st.caption("YouTube to Blog Post pipeline")
    st.divider()
    st.markdown("### Pipeline")
    st.markdown("""
    - ✅ **Agent 1** — Transcript
    - ✅ **Agent 2** — Blog Post
    """)
    st.divider()
    api_status = "✅ Set" if os.environ.get("GROQ_API_KEY") else "❌ Missing"
    st.caption(f"Groq API Key: {api_status}")
    st.caption("Groq Whisper + Llama 3.3 70B")
    st.divider()
    st.markdown("### Transcript Source")
    st.caption("⚡ Captions API — instant, no download")
    st.caption("🎙️ Audio download — fallback for videos without captions")

st.title("📝 YouTube to Blog Post")
st.caption("Paste a YouTube URL and get a publish-ready blog post.")

# ── Input ─────────────────────────────────────────────────────────────────────
url = st.text_input(
    label="YouTube URL",
    placeholder="https://www.youtube.com/watch?v=...",
    help="Any public YouTube video with speech will work.",
)

generate = st.button("Generate Blog Post", type="primary", use_container_width=True)

# ── API key check ─────────────────────────────────────────────────────────────
if not os.environ.get("GROQ_API_KEY"):
    st.warning("GROQ_API_KEY is not set. Add it in Railway → Variables.")

# ── Pipeline ──────────────────────────────────────────────────────────────────
if generate:
    if not url.strip():
        st.warning("Please paste a YouTube URL first.")
        st.stop()

    if not os.environ.get("GROQ_API_KEY"):
        st.error("GROQ_API_KEY is not set. Add it in Railway Variables.")
        st.stop()

    # ── Fetch video metadata and check database ────────────────────────────
    video_id   = ""
    title      = ""
    cached_row = None

    try:
        meta     = TranscriptAgent()._get_video_info(url.strip())
        video_id = meta["video_id"]
        title    = meta["title"]
    except Exception:
        pass  # metadata fetch failed; proceed without cache check

    if video_id:
        try:
            cached_row = get_video_by_id(video_id)
        except Exception as e:
            st.warning(f"DB lookup failed: {e}")
            cached_row = None

    # ── Cache hit: load from database ─────────────────────────────────────
    if cached_row:
        st.session_state["display"] = {
            "blog_post":         cached_row["blog_post"],
            "transcript":        cached_row["transcript"],
            "video_id":          cached_row["video_id"],
            "title":             cached_row.get("title", title),
            "transcript_source": "",
            "source":            "cached",
        }

    # ── Cache miss: run full pipeline ──────────────────────────────────────
    else:
        log_lines = []

        def append_log(msg: str):
            log_lines.append(msg)
            log_area.code("\n".join(log_lines), language=None)

        try:
            with st.status("Running pipeline...", expanded=True) as status:
                log_area = st.empty()

                append_log("Initialising agents...")
                orch = Orchestrator()

                original_print = builtins.print

                def captured_print(*args, **kwargs):
                    msg = " ".join(str(a) for a in args)
                    if msg.startswith("[Transcript") or msg.startswith("[BlogPost") or msg.startswith("[Orchestrator"):
                        append_log(msg)
                    original_print(*args, **kwargs)

                builtins.print = captured_print
                try:
                    result = orch.run(url.strip())
                finally:
                    builtins.print = original_print

                status.update(label="Done!", state="complete", expanded=False)

        except Exception as e:
            st.error(f"Something went wrong: {e}")
            st.stop()

        # Save to database (only when a valid video_id was returned)
        if result["video_id"]:
            try:
                save_video(
                    video_id   = result["video_id"],
                    url        = url.strip(),
                    title      = result["title"],
                    transcript = result["transcript"],
                    blog_post  = result["blog_post"],
                    logs       = "",
                )
            except Exception as e:
                st.warning(f"DB save failed: {e}")

        st.session_state["display"] = {
            "blog_post":         result["blog_post"],
            "transcript":        result["transcript"],
            "video_id":          result["video_id"],
            "title":             result["title"],
            "transcript_source": result.get("transcript_source", "audio_download"),
            "source":            "new",
        }

# ── Output (persists across re-runs via session state) ────────────────────────
if "display" in st.session_state:
    d = st.session_state["display"]

    if d["source"] == "cached":
        st.info("⚡ This video was already processed. Loaded from database instantly.")
    elif d["source"] == "new":
        st.success("✅ Blog post generated and saved to database.")

    st.divider()

    docx_bytes = blog_post_to_docx(d["blog_post"])

    # Cached and regenerated results include the Regenerate tab
    if d["source"] in ("cached", "regenerated"):
        tab1, tab2, tab3 = st.tabs(["Blog Post", "Raw Transcript", "Regenerate"])
    else:
        tab1, tab2       = st.tabs(["Blog Post", "Raw Transcript"])
        tab3             = None

    with tab1:
        st.markdown(d["blog_post"])
        st.download_button(
            label="Download as Word Document",
            data=docx_bytes,
            file_name="blog_post.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with tab2:
        if d["source"] == "cached":
            source_badge = "📦 Loaded from database"
        elif d["transcript_source"] == "captions_api":
            source_badge = "⚡ Via captions"
        else:
            source_badge = "🎙️ Via audio transcription"
        st.caption(source_badge)
        st.text_area(
            label="Transcript",
            value=d["transcript"],
            height=400,
            disabled=True,
        )
        st.caption(f"{len(d['transcript'].split())} words · {len(d['transcript'])} characters")

    if tab3 is not None:
        with tab3:
            st.caption("Run the AI writer again on the stored transcript to get a fresh blog post.")
            if st.button(
                "Generate a new blog post",
                key="regen_btn",
                type="primary",
                use_container_width=True,
            ):
                with st.spinner("Writing new blog post..."):
                    try:
                        blog_agent    = BlogPostAgent()
                        new_blog_post = blog_agent.run(d["transcript"])
                        try:
                            update_blog_post(d["video_id"], new_blog_post)
                        except Exception:
                            pass  # DB unavailable; still show the refreshed post
                        st.session_state["display"]["blog_post"] = new_blog_post
                        st.session_state["display"]["source"]    = "regenerated"
                        st.rerun()
                    except Exception as e:
                        st.error(f"Regeneration failed: {e}")
